import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re  # Add this line with your other imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor

def format_number_with_commas(value):
    """Format numerical values with commas (e.g., 5000 -> 5,000)"""
    try:
        # Handle both integers and floats
        if isinstance(value, (int, float)):
            return "{:,.2f}".format(value) if isinstance(value, float) else "{:,}".format(value)
        # Handle string representations of numbers
        elif isinstance(value, str) and value.replace('.', '').replace(',', '').isdigit():
            # Remove existing commas if present
            clean_value = value.replace(',', '')
            if '.' in clean_value:
                return "{:,.2f}".format(float(clean_value))
            else:
                return "{:,}".format(int(clean_value))
        return value
    except:
        return value  # Return original if formatting fails

def set_page_width(document):
    """Set the page width to accommodate wider content."""
    section = document.sections[0]
    section.page_width = Inches(8.3)  # Standard page width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def add_section_header(document, title):
    """Adds a formatted section header to the document."""
    document.add_paragraph()  # Add an empty line before the header
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True  # added underline
    run.font.name = 'Times New Roman'

def add_bold_subheading(document, text):
    """Adds a bold subheading within the document content."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'

def create_table(document, headers, data, include_total=False, total_column_index=None):
    """Creates a formatted table with given headers and data with auto-fit columns,
         optionally including a total row."""
    if not headers or not data:
        document.add_paragraph("No data available for this table.")
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True  # Enable auto-fit

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for item in data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            value = item.get(header, '')
            if header == 'Description':  # Check if it's the 'Description' column
                if isinstance(value, list):
                    p = row_cells[i].add_paragraph()
                    for item_desc in value:
                        p.add_run(f"• {item_desc}\n")
                    p.runs[0].font.name = 'Times New Roman'
                else:
                    row_cells[i].text = str(value)
            else:
                formatted_value = format_number_with_commas(value)
                row_cells[i].text = str(formatted_value)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'

    # Add total row if required
    if include_total and total_column_index is not None:
        try:
            total = sum(float(str(row[total_column_index]).replace(',', '')) for row in data)
            total_row_cells = table.add_row().cells
            total_row_cells[0].text = "Total"
            total_row_cells[total_column_index].text = format_number_with_commas(total)

            # Format the total row
            for cell in total_row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Times New Roman'
        except (IndexError, ValueError):
            print("Warning: Could not calculate or format total for the table.")

def create_table_with_subcolumns(document, main_headers, data, subheaders=None):
    """Creates a table with optional subcolumns and auto-fitting."""
    if not main_headers or not data:
        document.add_paragraph("No data available for this table.")
        return

    if subheaders and len(main_headers) != len(subheaders):
        raise ValueError("Length of main headers must match length of subheaders.")

    num_cols = sum(len(subs) if subs else 1 for subs in subheaders) if subheaders else len(main_headers)
    table = document.add_table(rows=1 + (1 if any(subheaders) else 0), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True  # Enable autofit

    hdr_cells = table.rows[0].cells
    col_index = 0
    for i, header in enumerate(main_headers):
        if subheaders and subheaders[i]:
            hdr_cells[col_index].text = header
            hdr_cells[col_index].paragraphs[0].runs[0].bold = True
            hdr_cells[col_index].paragraphs[0].runs[0].font.name = 'Times New Roman'
            hdr_cells[col_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hdr_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[col_index].merge(table.rows[0].cells[col_index + len(subheaders[i]) - 1])
            col_index += len(subheaders[i])
        else:
            hdr_cells[col_index].text = header
            hdr_cells[col_index].paragraphs[0].runs[0].bold = True
            hdr_cells[col_index].paragraphs[0].runs[0].font.name = 'Times New Roman'
            hdr_cells[col_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hdr_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            col_index += 1
    if any(subheaders):  # Add this condition
        sub_hdr_cells = table.add_row().cells
        col_index = 0
        for subs in subheaders:
            if subs:
                for sub in subs:
                    sub_hdr_cells[col_index].text = sub
                    sub_hdr_cells[col_index].paragraphs[0].runs[0].bold = True
                    sub_hdr_cells[col_index].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    sub_hdr_cells[col_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    sub_hdr_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    col_index += 1
            else:
                col_index += 1

    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        cell_index = 0
        for i, header_name in enumerate(main_headers):  # changed header to header_name
            value = row_data.get(header_name)
            if subheaders and subheaders[i]:
                if isinstance(value, dict):
                    for sub_header in subheaders[i]:
                        formatted_value = format_number_with_commas(value.get(sub_header, ''))
                        row_cells[cell_index].text = str(formatted_value)
                        row_cells[cell_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in row_cells[cell_index].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                        cell_index += 1
                else:
                    row_cells[cell_index].text = str(format_number_with_commas(value))
                    row_cells[cell_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in row_cells[cell_index].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                    cell_index += 1
            else:
                if header_name == 'Description':
                    if isinstance(value, list):
                        p = row_cells[cell_index].add_paragraph()
                        for item_desc in value:
                            p.add_run(f"• {item_desc}\n")
                        p.runs[0].font.name = 'Times New Roman'
                    else:
                        row_cells[cell_index].text = str(value)
                else:
                    formatted_value = format_number_with_commas(value)
                    row_cells[cell_index].text = str(formatted_value)
                row_cells[cell_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in row_cells[cell_index].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                cell_index += 1

def process_text_with_subheadings(document, text):
    """Processes text to identify and format subheadings (text followed by colons) as bold, all on one line."""
    if not text:
        return

    # Handle the case where the text is a dictionary
    if isinstance(text, dict):
        # Check for the specific 'supply_and_demand_analysis' key
        if 'Supply_and_Demand_Analysis' in text:
            format_supply_demand_analysis(document, text['Supply_and_Demand_Analysis'])
            return  # Stop processing here, as we've handled it
        else:
            for key, value in text.items():
                process_text_with_subheadings(document, f"{key}: {value}")
            return

    # Handle the case where the text is a list
    if isinstance(text, list):
        for item in text:
            process_text_with_subheadings(document, item)
        return

    # Remove any unwanted stars or markdown-like formatting
    clean_text = text.replace('**', '').replace('*', '').strip()

    lines = clean_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if ':' in line and len(line.split(':')[0].strip()) < 30:  # Heuristic for subheadings
            parts = line.split(':', 1)
            p = document.add_paragraph()
            run1 = p.add_run(parts[0] + ': ')
            run1.bold = True
            run1.font.name = 'Times New Roman'
            run2 = p.add_run(parts[1].strip())
            run2.font.name = 'Times New Roman'
        elif line.startswith('- '):  # Handle bullet points
            p = document.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:].strip())
            run.font.name = 'Times New Roman'
        else:
            p = document.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Times New Roman'



def process_ict_requirements(document, ict_text):
    """Processes the text for the ICT Requirements section, making subheadings bold."""
    if not ict_text:
        return

    # Handle dictionary input
    if isinstance(ict_text, dict):
        for category, details in ict_text.items():
            # Add category as bold subheading
            p = document.add_paragraph()
            run = p.add_run(f"{category.replace('_', ' ').title()}:")
            run.bold = True
            run.font.name = 'Times New Roman'
            
            # Handle nested dictionaries
            if isinstance(details, dict):
                for subcategory, spec in details.items():
                    if isinstance(spec, dict):
                        # For multi-level nested specs (like hardware specifications)
                        p = document.add_paragraph()
                        run = p.add_run(f"  {subcategory.replace('_', ' ').title()}:")
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        
                        for key, value in spec.items():
                            document.add_paragraph(f"    • {key.replace('_', ' ').title()}: {value}", style='List Bullet')
                    else:
                        document.add_paragraph(f"  • {subcategory.replace('_', ' ').title()}: {spec}", style='List Bullet')
            elif isinstance(details, list):
                # Handle lists (like software requirements)
                for item in details:
                    document.add_paragraph(f"  • {item}", style='List Bullet')
            else:
                document.add_paragraph(f"  • {details}", style='List Bullet')
        return

    # Original text processing for string input
    clean_text = ict_text.replace('**', '').replace('*', '').strip()
    lines = clean_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:].strip())
            run.font.name = 'Times New Roman'
        else:
            if line.startswith('·\xa0\xa0\xa0'):
                potential_subheading = line[4:].strip()
                if ':' in potential_subheading and len(potential_subheading.split(':')[0].strip()) < 30:
                    parts = potential_subheading.split(':', 1)
                    p = document.add_paragraph()
                    run1 = p.add_run(parts[0] + ': ')
                    run1.bold = True
                    run1.font.name = 'Times New Roman'
                    run2 = p.add_run(parts[1].strip())
                    run2.font.name = 'Times New Roman'
                else:
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    run.font.name = 'Times New Roman'
            elif ':' in line and len(line.split(':')[0].strip()) < 30:
                parts = line.split(':', 1)
                p = document.add_paragraph()
                run1 = p.add_run(parts[0] + ': ')
                run1.bold = True
                run1.font.name = 'Times New Roman'
                run2 = p.add_run(parts[1].strip())
                run2.font.name = 'Times New Roman'
            else:
                p = document.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Times New Roman'


def format_supply_demand_analysis(document, data):
    """Formats the Supply and Demand Analysis section."""
    add_bold_subheading(document, "Supply and Demand Analysis")  # Overall section title

    if isinstance(data, str):
        # If it's a string, just process it as text
        process_text_with_subheadings(document, data)
        return

    if not isinstance(data, dict):
        # If it's not a dictionary, convert to string and process
        process_text_with_subheadings(document, str(data))
        return

    # Process dictionary structure
    for section_title, section_content in data.items():
        # Clean up section title (remove numbers if present)
        clean_title = re.sub(r'^\d+\.\s*', '', section_title).title()
        add_bold_subheading(document, clean_title)

        if isinstance(section_content, dict):
            for subsection, content in section_content.items():
                p = document.add_paragraph()
                run = p.add_run(f"  {subsection.replace('_', ' ').title()}: ")
                run.bold = True
                run.font.name = 'Times New Roman'

                if isinstance(content, dict):
                    # Handle nested dictionaries
                    for key, value in content.items():
                        document.add_paragraph(f"    • {key.replace('_', ' ').title()}: {value}", style='List Bullet')
                elif isinstance(content, list):
                    # Handle lists
                    for item in content:
                        document.add_paragraph(f"    • {item}", style='List Bullet')
                else:
                    # Handle simple values
                    p.add_run(str(content))
        elif isinstance(section_content, list):
            # Handle top-level lists
            for item in section_content:
                document.add_paragraph(f"• {item}", style='List Bullet')
        else:
            # Handle simple values
            p = document.add_paragraph()
            p.add_run(str(section_content))

def create_project_document_from_json(json_data, output_docx_path):
    """Creates a Word document with information from the provided JSON data."""
    doc = Document()

    # Add "PC-1 FORM" heading at the top center
    pc1_heading = doc.add_paragraph()
    pc1_heading_run = pc1_heading.add_run("PC-1 FORM\nGOVERNMENT OF PAKISTAN\nPLANNING COMMISSION")
    pc1_heading_run.bold = True
    pc1_heading_run.font.size = Pt(16)  # Adjust font size as needed
    pc1_heading_run.font.name = 'Times New Roman'
    pc1_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Add an empty line after the heading

    # Set page width and margins
    set_page_width(doc)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    project_data = json_data

    # Add basic information sections
    add_section_header(doc, "1. Project Overview")
    p = doc.add_paragraph()
    run1 = p.add_run("Project Name: ")
    run1.font.name = 'Times New Roman'
    run1.bold = True
    run2 = p.add_run(f"{project_data.get('projectName', 'N/A')}")
    run2.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    run1 = p.add_run("District: ")
    run1.font.name = 'Times New Roman'
    run1.bold = True
    run2 = p.add_run(f"{project_data.get('districtName', 'N/A')}")
    run2.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    run1 = p.add_run("Sector: ")
    run1.font.name = 'Times New Roman'
    run1.bold = True
    run2 = p.add_run(f"{project_data.get('sector', 'N/A')}")
    run2.font.name = 'Times New Roman'

    add_section_header(doc, "2. Agency Information")
    agency_info = [
        {"Sponsoring Agency": project_data.get("sponsAgency", "N/A"),
         "Operating Agency": project_data.get("opAgency", "N/A"),
         "Executing Agency": project_data.get("exeAgency", "N/A"),
         "Maintenance Agency": project_data.get("maintAgency", "N/A")},
        {"Is Provincial": project_data.get("isProvincial", "N/A"),
         "Federal Ministry": project_data.get("federalMinistry", "N/A")}
    ]
    create_table(doc, list(agency_info[0].keys()), [agency_info[0]])
    create_table(doc, list(agency_info[1].keys()), [agency_info[1]])

    add_section_header(doc, "3. Project Timeline & Budget")
    timeline_budget = [
        {"Budget": project_data.get("budget", "N/A"),
         "Duration (months)": project_data.get("duration", "N/A"),
         "Start Date": project_data.get("startDate", "N/A"),
         "End Date": project_data.get("endDate", "N/A")}
    ]
    create_table(doc, list(timeline_budget[0].keys()), timeline_budget)

    add_section_header(doc, "4. Project Details")
    process_text_with_subheadings(doc, f"Scope: {project_data.get('scope', 'N/A')}")
    process_text_with_subheadings(doc, f"Location: {project_data.get('location', 'N/A')}")
    process_text_with_subheadings(doc, f"Feasibility Study Completed: {project_data.get('feasibilityStudy', 'N/A')}")
    process_text_with_subheadings(doc, f"Design Finalized: {project_data.get('designFinalized', 'N/A')}")
    process_text_with_subheadings(doc, f"Technology: {project_data.get('technology', 'N/A')}")
    process_text_with_subheadings(doc, f"Capacity: {project_data.get('capacity', 'N/A')}")
    process_text_with_subheadings(doc, f"Phases: {project_data.get('phases', 'N/A')}")

    add_section_header(doc, "5. Risks, Deliverables & Stakeholders")
    process_text_with_subheadings(doc, f"Risks: {project_data.get('risks', 'N/A')}")
    process_text_with_subheadings(doc, f"Deliverables: {project_data.get('deliverables', 'N/A')}")
    process_text_with_subheadings(doc, f"Stakeholders: {project_data.get('stakeholders', 'N/A')}")

    add_section_header(doc, "6. Monitoring & Sustainability")
    process_text_with_subheadings(doc, f"Monitoring Plan: {project_data.get('monitoringPlan', 'N/A')}")
    process_text_with_subheadings(doc, f"Sustainability Measures: {project_data.get('sustainabilityMeasures', 'N/A')}")

    add_section_header(doc, "7. Objectives")
    process_text_with_subheadings(doc, project_data.get("Objectives", "N/A"))

    add_section_header(doc, "8. ICT Requirements")
    ict_reqs = project_data.get("ICT-Reqs", "N/A")
    if isinstance(ict_reqs, dict):
        process_ict_requirements(doc, ict_reqs)
    elif isinstance(ict_reqs, str):
        ict_reqs = ict_reqs.replace('**', '').replace('*', '')
        # Try to parse as JSON if it's a string that might contain JSON
        try:
            ict_data = json.loads(ict_reqs)
            process_ict_requirements(doc, ict_data)
        except json.JSONDecodeError:
            # Process as plain text if not JSON
            lines = ict_reqs.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    process_ict_requirements(doc, line)
    else:
        process_ict_requirements(doc, str(ict_reqs))

    add_section_header(doc, "9. Supply and Demand Analysis")
    supply_demand = project_data.get("Supply and Demand", {})
    if isinstance(supply_demand, str) and supply_demand.startswith("Download\n"):
        try:
            supply_demand = json.loads(supply_demand.replace("Download\n", ""))
        except json.JSONDecodeError:
            pass  # Keep as string if can't parse

    format_supply_demand_analysis(doc, supply_demand)

    # --- Process Capital Cost Estimates ---
    add_section_header(doc, "10. Capital Cost Estimates")
    capital_cost_estimates = project_data.get("capitalCostEstimates", "")

    if isinstance(capital_cost_estimates, str) and capital_cost_estimates.startswith("Download\n"):
        capital_cost_data = json.loads(capital_cost_estimates.replace("Download\n", ""))
    elif isinstance(capital_cost_estimates, dict):
        capital_cost_data = capital_cost_estimates
    else:
        capital_cost_data = {}

    if isinstance(capital_cost_data, dict) and 'capitalCost' in capital_cost_data:
        for cost_section in capital_cost_data['capitalCost']:
            add_bold_subheading(doc, cost_section.get('name', ''))
            process_text_with_subheadings(doc, cost_section.get('description', ''))
            if 'data' in cost_section and cost_section['data']:
                first_data_item = cost_section['data'][0]
                has_subheaders = any(isinstance(value, dict) for value in first_data_item.values())

                if has_subheaders:
                    main_headers = []
                    subheaders = []
                    for key, value in first_data_item.items():
                        if isinstance(value, dict):
                            main_headers.append(key)
                            subheaders.append(list(value.keys()))
                        else:
                            main_headers.append(key)
                            subheaders.append(None)
                    create_table_with_subcolumns(doc, main_headers, cost_section['data'],
                                                subheaders=subheaders)
                else:
                    headers = list(cost_section['data'][0].keys()) if cost_section['data'] else []
                    create_table(doc, headers, cost_section['data'])

    # --- Process Maintenance Costs ---
    add_section_header(doc, "11. Maintenance Costs")

    maintenance_costs = project_data.get("maintenanceCosts", "")
    if isinstance(maintenance_costs, str) and maintenance_costs.startswith("Download\n"):
        try:
            maintenance_costs_data = json.loads(maintenance_costs.replace("Download\n", ""))
        except json.JSONDecodeError:
            maintenance_costs_data = {}
    elif isinstance(maintenance_costs, dict):
        maintenance_costs_data = maintenance_costs
    else:
        maintenance_costs_data = {}

    if isinstance(maintenance_costs_data, dict):
        if 'financialPlan' in maintenance_costs_data and maintenance_costs_data['financialPlan']:
            add_bold_subheading(doc, "Financial Plan:")
            table_data = []
            for item in maintenance_costs_data['financialPlan']:
                for key, value in item.items():
                    if "Year" in key or "Total Cost" in key:
                        year_part = key
                        cost_part = value
                        table_data.append({"Year": year_part, "Amount (Rs. in million)": cost_part})
            if table_data:
                headers = ["Year", "Amount (Rs. in million)"]
                create_table(doc, headers, table_data)

        elif 'operations' in maintenance_costs_data and maintenance_costs_data['operations']:
            add_bold_subheading(doc, "Operations Costs:")
            operations_data = [
                {"Description": d, "Amount (Rs. in million)": a}
                for d, a in zip(
                    maintenance_costs_data['operations'].get('description', []),
                    maintenance_costs_data['operations'].get('Amount', [])
                )
            ]
            if operations_data:
                headers = ["Description", "Amount (Rs. in million)"]
                create_table(doc, headers, operations_data)

    add_section_header(doc, "12. Benefits")

    benefits = project_data.get("benefits", "")
    if isinstance(benefits, str) and benefits.startswith("Download\n"):
        try:
            benefits_data = json.loads(benefits.replace("Download\n", ""))
        except json.JSONDecodeError:
            benefits_data = {}
    elif isinstance(benefits, dict):
        benefits_data = benefits
    else:
        benefits_data = {}

    if isinstance(benefits_data, dict) and 'project_components' in benefits_data:
        add_bold_subheading(doc, "Project Components:")
        headers = [
            "S.No.", "Input", "Component", "Units",
            "Year 1 Amount", "Year 1 Division",
            "Year 2 Amount", "Year 2 Division",
            "Year 3 Amount", "Year 3 Division",
            "Baseline Indicator", "Post Completion Targets",
            "Key Benefits"
        ]
        table_data = []
        for item in benefits_data['project_components']:
            table_data.append({
                "S.No.": item.get('serial_number', ''),
                "Input": item.get('input', ''),
                "Component": item.get('outcome', {}).get('component_name', ''),
                "Units": item.get('outcome', {}).get('units', ''),
                "Year 1 Amount": item.get('year_wise_phasing', {}).get('year_1', {}).get('amount', ''),
                "Year 1 Division": item.get('year_wise_phasing', {}).get('year_1', {}).get(
                    'division_of_total_items', ''),
                "Year 2 Amount": item.get('year_wise_phasing', {}).get('year_2', {}).get('amount', ''),
                "Year 2 Division": item.get('year_wise_phasing', {}).get('year_2', {}).get(
                    'division_of_total_items', ''),
                "Year 3 Amount": item.get('year_wise_phasing', {}).get('year_3', {}).get('amount', ''),
                "Year 3 Division": item.get('year_wise_phasing', {}).get('year_3', {}).get(
                    'division_of_total_items', ''),
                "Baseline Indicator": item.get('outcome_metrics', {}).get('baseline_indicator', ''),
                "Post Completion Targets": item.get('targeted_impact', {}).get(
                    'post_completion_targets', ''),
                "Key Benefits": item.get('impact_details', {}).get('key_benefits', '')
            })
        create_table(doc, headers, table_data)

    add_section_header(doc, "13. Financial Plan Table")
    financial_plan_table = project_data.get("financialPlanTable", {})
    if isinstance(financial_plan_table, dict):
        financial_plan_data = financial_plan_table.get("financialPlan", [])
        if financial_plan_data:
            headers = list(financial_plan_data[0].keys()) if financial_plan_data else []
            create_table(doc, headers, financial_plan_data)
        else:
            process_text_with_subheadings(doc, json.dumps(financial_plan_table, indent=2))
    elif isinstance(financial_plan_table, str):
        process_text_with_subheadings(doc, financial_plan_table)
    else:
        process_text_with_subheadings(doc, str(financial_plan_table))

    add_section_header(doc, "14. Management Structure and Manpower")
    process_text_with_subheadings(doc, project_data.get("managementStructure", "N/A"))

    add_section_header(doc, "15. Additional Projects/Decisions")
    process_text_with_subheadings(doc, project_data.get("additionalProjects", "N/A"))

    add_section_header(doc, "16. Certification")
    process_text_with_subheadings(doc,
                                    "Certified that the project proposal has been prepared...")  # Add full certification text as needed
    process_text_with_subheadings(doc, f"Prepared by: {project_data.get('prepared_by', 'N/A')}")
    process_text_with_subheadings(doc, f"Checked by: {project_data.get('checked_by', 'N/A')}")
    process_text_with_subheadings(doc, f"Approved by: {project_data.get('approved_by', 'N/A')}")

    # Save the document
    doc.save(output_docx_path)
    print(f"Document generated successfully at: {output_docx_path}")
